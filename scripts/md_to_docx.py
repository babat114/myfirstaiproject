"""
Convert Markdown files to properly formatted .docx documents.
Handles: headings, tables, code blocks, lists, bold/inline code, horizontal rules.

Usage:
  python scripts/md_to_docx.py 需求规格说明书.md
  python scripts/md_to_docx.py PRD.md
"""
import re
import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


class MarkdownToDocx:
    """Line-by-line markdown → docx converter with table support."""

    def __init__(self, output_path, font_name='Microsoft YaHei', font_size=11):
        self.doc = Document()
        self.output_path = output_path
        self.font_name = font_name
        self.font_size = font_size

        # Configure default style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        # Configure heading styles
        for level in range(1, 5):
            h_style = self.doc.styles[f'Heading {level}']
            h_font = h_style.font
            h_font.name = font_name
            h_style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if level == 1:
                h_font.size = Pt(22)
                h_font.bold = True
                h_font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
            elif level == 2:
                h_font.size = Pt(16)
                h_font.bold = True
                h_font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
            elif level == 3:
                h_font.size = Pt(13)
                h_font.bold = True
            elif level == 4:
                h_font.size = Pt(11.5)
                h_font.bold = True

        # Page margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    def _shade_cell(self, cell, color):
        """Set cell background color."""
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _set_cell_borders(self, cell, **kwargs):
        """Set cell borders."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        for edge, val in kwargs.items():
            edge_el = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{val.get("sz",4)}" w:space="0" w:color="{val.get("color","999999")}"/>')
            tcBorders.append(edge_el)
        tcPr.append(tcBorders)

    def _add_formatted_paragraph(self, text, style='Normal'):
        """Add paragraph with inline formatting (bold, code, links)."""
        para = self.doc.add_paragraph(style=style)

        # Split text into segments: **bold**, `code`, [link](url), plain
        segments = re.split(r'(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))', text)

        for seg in segments:
            if not seg:
                continue
            if seg.startswith('**') and seg.endswith('**'):
                run = para.add_run(seg[2:-2])
                run.bold = True
            elif seg.startswith('`') and seg.endswith('`'):
                run = para.add_run(seg[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(self.font_size - 1)
                # Gray background for inline code
                run.font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)
            elif seg.startswith('[') and '](' in seg:
                m = re.match(r'\[(.*?)\]\((.*?)\)', seg)
                if m:
                    run = para.add_run(m.group(1))
                    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
                    run.underline = True
            else:
                para.add_run(seg)

        return para

    def _add_code_block(self, lines):
        """Add a code block (lines between ``` markers)."""
        code_text = '\n'.join(lines)
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

        # Add shading to code block
        pPr = para._p.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F4F4"/>')
        pPr.append(shading)

        run = para.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        return para

    def _add_table(self, header_row, data_rows):
        """Add a formatted table."""
        if not header_row:
            return

        num_cols = len(header_row)
        num_rows = 1 + len(data_rows)
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for j, cell_text in enumerate(header_row):
            cell = table.rows[0].cells[j]
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(cell_text.strip())
            run.bold = True
            run.font.size = Pt(self.font_size - 1)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._shade_cell(cell, '2c3e50')

        # Data rows
        for i, row in enumerate(data_rows):
            for j, cell_text in enumerate(row):
                if j < num_cols:
                    cell = table.rows[i + 1].cells[j]
                    cell.text = ''
                    para = cell.paragraphs[0]
                    # Handle bold in cells
                    text = cell_text.strip()
                    segs = re.split(r'(\*\*.*?\*\*)', text)
                    for seg in segs:
                        if seg.startswith('**') and seg.endswith('**'):
                            run = para.add_run(seg[2:-2])
                            run.bold = True
                        else:
                            para.add_run(seg)
                    para.style.font.size = Pt(self.font_size - 1)
                    # Alternate row shading
                    if i % 2 == 1:
                        self._shade_cell(cell, 'f2f6fc')

        # Add spacing after table
        self.doc.add_paragraph()

    def _is_table_separator(self, line):
        """Check if a line is a markdown table separator: |---|---|"""
        return bool(re.match(r'^\|[\s\-:|]+\|$', line))

    def _parse_table_row(self, line):
        """Parse a markdown table row into cells."""
        # Split by | and trim whitespace
        cells = [c.strip() for c in line.split('|')]
        # Remove empty first/last from leading/trailing |
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        return cells

    def convert(self, md_path):
        """Main conversion logic."""
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        i = 0
        in_code_block = False
        code_lines = []
        in_table = False
        table_header = []
        table_data = []
        in_list = False

        while i < len(lines):
            line = lines[i].rstrip()

            # --- Code blocks ---
            if line.startswith('```'):
                if in_code_block:
                    self._add_code_block(code_lines)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                    code_lines = []
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # --- Empty lines: flush pending structures ---
            if not line.strip():
                if in_table:
                    self._add_table(table_header, table_data)
                    table_header = []
                    table_data = []
                    in_table = False
                if in_list:
                    in_list = False
                self.doc.add_paragraph()  # spacer
                i += 1
                continue

            # --- Table detection ---
            if '|' in line and self._is_table_separator(line):
                # Previous line was the header
                if table_header:
                    in_table = True
                i += 1
                continue

            if in_table:
                row = self._parse_table_row(line)
                if row:
                    table_data.append(row)
                i += 1
                continue

            if '|' in line and not in_table:
                # Potential table header — peek ahead
                if i + 1 < len(lines) and self._is_table_separator(lines[i + 1].rstrip()):
                    table_header = self._parse_table_row(line)
                    i += 1  # skip header, separator will be caught next iteration
                    continue

            # --- Headings ---
            h_match = re.match(r'^(#{1,4})\s+(.+)', line)
            if h_match:
                level = len(h_match.group(1))
                text = h_match.group(2)
                # Remove markdown links from heading text
                text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
                self.doc.add_heading(text, level=level)
                i += 1
                continue

            # --- Horizontal rule ---
            if re.match(r'^[-*_]{3,}$', line.strip()):
                para = self.doc.add_paragraph()
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(12)
                # Add a thin horizontal line
                pPr = para._p.get_or_add_pPr()
                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
                    f'</w:pBdr>'
                )
                pPr.append(pBdr)
                i += 1
                continue

            # --- Lists ---
            list_match = re.match(r'^(\s*)([-*]\s+|\d+\.\s+)(.+)', line)
            if list_match:
                indent = len(list_match.group(1))
                text = list_match.group(3)
                para = self._add_formatted_paragraph(text)
                para.paragraph_format.left_indent = Cm(1.5 + indent * 0.5)
                if list_match.group(2).startswith(('-', '*')):
                    para.style = self.doc.styles['List Bullet']
                else:
                    para.style = self.doc.styles['List Number']
                i += 1
                continue

            # --- Blockquote ---
            if line.startswith('> '):
                text = line[2:].strip()
                if text.startswith('**') and text.endswith('**'):
                    # Blockquote heading
                    para = self._add_formatted_paragraph(text)
                    para.paragraph_format.left_indent = Cm(1)
                else:
                    para = self._add_formatted_paragraph(text)
                    para.paragraph_format.left_indent = Cm(1)
                    para.runs[0].font.italic = True if para.runs else False
                i += 1
                continue

            # --- Regular paragraph ---
            self._add_formatted_paragraph(line)
            i += 1

        # Flush any remaining table
        if in_table and table_header:
            self._add_table(table_header, table_data)

        # Save
        self.doc.save(self.output_path)
        print(f'  Created: {self.output_path}')


def main():
    if len(sys.argv) < 2:
        print('Usage: python scripts/md_to_docx.py <markdown_file>')
        print('  e.g.: python scripts/md_to_docx.py 需求规格说明书.md')
        sys.exit(1)

    md_path = sys.argv[1]
    if not os.path.exists(md_path):
        print(f'Error: file not found: {md_path}')
        sys.exit(1)

    # Output path: same name with .docx extension
    base = os.path.splitext(md_path)[0]
    output = f'{base}.docx'

    print(f'Converting: {md_path} → {output}')
    converter = MarkdownToDocx(output)
    converter.convert(md_path)
    print('Done.')


if __name__ == '__main__':
    main()
