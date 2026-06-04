"""
============================================
PRD Markdown → Word 转换脚本
将 PRD.md 转换为格式化的 Word 文档
运行: python scripts/convert_prd_to_docx.py
============================================
"""
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_code_block(doc, code_text, font_size=9):
    """添加代码块 (灰色背景)"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.style = doc.styles['Code']
        run = p.add_run(line)
        run.font.size = Pt(font_size)
        run.font.name = 'Consolas'


def add_table_row(table, cells_data, is_header=False):
    """添加表格行"""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        if is_header:
            run.bold = True
            set_cell_shading(cell, '2B579A')
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return row


def convert_prd_to_docx(md_path, output_path):
    """将 PRD.md 转换为格式化的 Word 文档"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ---- 自定义样式 ----
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.35
    # 设置中文字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题样式
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = '微软雅黑'
        h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if i == 1:
            h_style.font.size = Pt(22)
            h_style.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        elif i == 2:
            h_style.font.size = Pt(16)
            h_style.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
        else:
            h_style.font.size = Pt(13)
            h_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 代码样式
    code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(2)
    code_style.paragraph_format.left_indent = Cm(0.5)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    shading.set(qn('w:val'), 'clear')
    code_style.paragraph_format.element.get_or_add_pPr().append(shading)

    # ---- 封面标题 ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(120)
    run = p.add_run('AI模型训练管理平台 v2.0')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('产品需求文档 (PRD) — 轻量增强版')
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2.font.name = '微软雅黑'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.space_after = Pt(60)
    run3 = p3.add_run('2026年6月  |  版本 2.0')
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ---- 解析 Markdown 内容 ----
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_headers = []
    table_rows = []
    table_obj = None

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # 空行
        if not line.strip():
            # 结束表格
            if in_table and table_obj is not None:
                in_table = False
                table_obj = None
            i += 1
            continue

        # 标题
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # 水平线
        if line.strip() == '---':
            doc.add_paragraph('─' * 60)
            i += 1
            continue

        # 表格
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]

            # 检查是否是分隔行
            if all(re.match(r'^[-:]+$', c) for c in cells):
                if not in_table:
                    # 创建表格
                    table_obj = doc.add_table(rows=1, cols=len(table_headers))
                    table_obj.style = 'Light Grid Accent 1'
                    table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
                    # 填充表头
                    hdr = table_obj.rows[0]
                    for j, h in enumerate(table_headers):
                        cell = hdr.cells[j]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        run = p.add_run(h)
                        run.bold = True
                        run.font.size = Pt(10)
                        set_cell_shading(cell, '2B579A')
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    in_table = True
                i += 1
                continue
            else:
                if in_table and table_obj is not None:
                    row = table_obj.add_row()
                    for j, c in enumerate(cells):
                        if j < len(row.cells):
                            row.cells[j].text = ''
                            p = row.cells[j].paragraphs[0]
                            run = p.add_run(c)
                            run.font.size = Pt(10)
                else:
                    table_headers = cells
                i += 1
                continue

        # ASCII 图表 — 跳过
        if line.strip().startswith('┌') or line.strip().startswith('│') or line.strip().startswith('└') or line.strip().startswith('├'):
            i += 1
            continue

        # 流程图箭头行 — 跳过
        if line.strip() in ('▼', '│', '──→', '───→') or set(line.strip()).issubset({'│', ' ', '▼', '├', '└', '►', '─', '┌', '┐', '┘', '└', ' '}):
            i += 1
            continue

        # 路径展示 — 特殊处理
        if line.strip() and ('app/' in line or 'experiments/' in line) and not any(kw in line for kw in ['import ', 'from ', 'class ', 'def ', ' = ', 'return']):
            # 可能是文件树展示
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(9)
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # 普通段落
        # 去除粗体标记
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', line.strip())
        # 去除行内代码
        text = re.sub(r'`([^`]+)`', r'\1', text)
        # 去除链接
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

        # 检查是否是列表项
        if re.match(r'^\d+\.\s', text):
            text = re.sub(r'^\d+\.\s', '', text)
            p = doc.add_paragraph(text, style='List Number')
        elif text.startswith('- '):
            text = text[2:]
            p = doc.add_paragraph(text, style='List Bullet')
        else:
            p = doc.add_paragraph(text)

        i += 1

    # ---- 页脚 ----
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = p_footer.add_run('— 文档结束 —')
    run_f.font.size = Pt(9)
    run_f.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run_f.italic = True

    # ---- 保存 ----
    doc.save(output_path)
    print(f'[OK] PRD converted to Word: {output_path}')


if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    md_path = os.path.join(base_dir, 'PRD.md')
    output_path = os.path.join(base_dir, 'PRD_AI_Platform_v2.0.docx')

    if not os.path.exists(md_path):
        print(f'[ERROR] PRD.md not found: {md_path}')
        sys.exit(1)

    convert_prd_to_docx(md_path, output_path)
