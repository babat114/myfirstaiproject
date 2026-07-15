#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate the training report (docx) for the AI Model Training Management Platform.

Combines two template files:
  1. Report template (封面 + 正文框架)
  2. Assessment form (考核表, appended after page break)
"""

from __future__ import annotations

import copy
import dataclasses
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from docx.shared import Cm, Pt, RGBColor
from lxml import etree


# ──────────────────────────────────────────────
#  Configuration (dataclass)
# ──────────────────────────────────────────────

@dataclasses.dataclass(frozen=True)
class StudentInfo:
    """Personal information for the report cover."""
    name: str = '李想'
    student_id: str = '2023081213'
    class_name: str = '计科232'
    teacher: str = '赵迪'
    year: str = '2025'


@dataclasses.dataclass(frozen=True)
class FileConfig:
    """File paths for templates and output."""
    template_dir: Path = Path(
        r'C:\Users\86180\xwechat_files\wxid_1a9j7sbt5m1m22_b42c\msg\file\2026-07'
    )
    template_report: Path = Path(
        r'C:\Users\86180\xwechat_files\wxid_1a9j7sbt5m1m22_b42c\msg\file\2026-07'
    ) / '校内专业实训(web软件开发方向)报告 - 格式模板.docx'
    template_form: Path = Path(
        r'C:\Users\86180\xwechat_files\wxid_1a9j7sbt5m1m22_b42c\msg\file\2026-07'
    ) / '1-2双面打印-《专业综合实训II(校企合作)》考核表-每人一份附在课程报告后面(1).docx'
    output_dir: Path = Path(r'C:\Users\86180\Desktop')
    output_name: str = '校内专业实训(web软件开发方向)报告 - 李想.docx'

    @property
    def output_path(self) -> Path:
        return self.output_dir / self.output_name


# ──────────────────────────────────────────────
#  Report content constants
# ──────────────────────────────────────────────

CH1_TITLE = '一、需求分析'
CH1_BODY = """
1.1 项目背景

随着人工智能技术的快速发展，机器学习模型的数据集管理与训练流程管理成为企业级AI应用的关键挑战。传统工作模式中，数据科学家需要手动管理数据集文件、编写训练脚本、记录实验结果，效率低下且难以复现。

本项目"AI模型训练管理平台"（myfirstaiproject）旨在构建一个基于B/S架构的全栈Web应用，实现数据集上传管理、模型训练任务调度、超参数自动调优、模型质量评估与推理服务的全生命周期管理。平台以"零新增第三方依赖、本地直接运行"为设计原则，基于Python标准库和Flask 3.1.0框架实现完整的MLOps工作流，无需Celery/Redis等额外中间件。

1.2 功能需求

系统围绕"数据集→模型训练→模型管理→模型推理"的核心链路，设计了以下七大功能模块：

（1）用户认证模块：支持用户名/邮箱注册与登录，三类角色权限管理（admin管理员/researcher研究员/viewer查看者），三种认证方式共存——Web端Session（Flask-Login）、API端JWT双Token认证（access_token 30分钟 + refresh_token 7天）和API Key兼容认证。

（2）数据集管理模块：支持CSV/JSON/Parquet/XLSX/TXT多种格式上传，自动解析列类型和统计摘要，11种分类（分类/回归/聚类/NLP/视觉/时序/生物/金融等），集成15个公开数据源自动采集，自持从uploading到ready的状态生命周期管理。

（3）模型注册表模块：支持8种模型类型和5种框架（scikit-learn 1.8.0、PyTorch 2.6.0+cu124 GPU版、TensorFlow/Keras 2.20.0、ONNX、Other）的模型注册与版本管理，记录accuracy/precision/recall/F1/R2/MSE等性能指标，提供排行榜与多模型对比功能。

（4）训练任务管理模块：基于ThreadPoolExecutor并发引擎（max_workers=2），支持任务创建/暂停/恢复/取消全生命周期管理，SSE（Server-Sent Events）实现500ms间隔实时进度推送，7种状态（queued→preparing→running↔paused→completed/failed/cancelled）流转。

（5）超参数调优模块：三种搜索策略——GridSearchCV穷举搜索（小参数空间）、RandomSearchCV随机采样（大参数空间）、AutoML自动遍历所有算法生成排名（未知最优算法时），聚类任务改用ParameterGrid + silhouette_score评估，TuningProgressTracker单例以500ms推送实时进度。

（6）模型预测与推理模块：交互式模型测试（文本/特征/文件/图像四种输入模式），实时NLP情感预测防抖（500ms debounce），批量预测与结果导出（CSV/JSON），Docker一键部署包生成，HuggingFace风格模型卡片展示。

（7）质量治理模块：7维模型质量自动诊断（训练精度/验证精度/泛化差距/过拟合指数/常数预测检测/特征重要性/残差分析），Tier 1/2/3三级分级管理，自动看门狗（once扫描/daemon守护/quick快速三种模式），过拟合和常数预测器自动检测与清理。

1.3 用户角色

系统定义了三类用户角色：管理员（admin）拥有系统全部管理权限，包括用户管理、角色分配和系统配置；研究员（researcher）可以使用全部功能，包括数据集上传、模型训练、超参数调优和模型推理；查看者（viewer）仅具有只读权限，可浏览数据集、模型和训练结果，适用于项目展示和审计场景。

1.4 非功能需求

系统需支持实时训练进度推送（500ms间隔）、并发训练任务管理（最多2个任务同时运行）、GPU自动检测与加速（NVIDIA RTX 4060 Laptop 8GB VRAM，CUDA 12.4）、Windows GBK编码环境兼容、375+自动化测试保证回归质量、SQLAlchemy 2.0 ORM统一数据访问层。
"""

CH2_TITLE = '二、系统设计'
CH2_BODY = """
2.1 技术架构

系统采用经典的四层架构设计：

表现层使用Flask Jinja2模板引擎 + Bootstrap 5.3响应式布局 + Vue 3（CDN） + Element Plus组件库 + Chart.js 4.4图表，构建现代化数据管理界面。Frontend v4.0新增Canvas粒子网络背景、滚动揭示动画和3D卡片倾斜效果。

路由层通过Flask Blueprint模块化组织，分为6个Web蓝图（auth/dashboard/datasets/models/training/comments）和7个API蓝图（v1认证/数据集/模型/训练/SSE流/用户/评论），统一通过parse_form_params()进行参数类型转换。

服务层封装认证服务、数据集服务、模型服务、训练服务、推理服务、超参数调优服务等核心业务逻辑，各服务间通过依赖注入解耦。

数据层采用MySQL 8.0数据库（开发环境SQLite内存数据库），4个核心ORM模型——User、Dataset、ModelRecord、TrainingJob，配合文件系统存储数据集和模型文件。

2.2 认证架构

系统实现了三种认证方式共存的灵活架构。Web端使用Flask-Login的Session认证，WTForms CSRF全局保护；API端使用PyJWT 2.10.1签发Bearer Token，支持token_version令牌版本控制（密码修改后旧令牌立即失效）；保留API Key（X-API-Key Header，SHA256哈希存储）实现向后兼容。三者通过统一的get_current_user()函数解析认证源，@api_login_required和@api_admin_required装饰器简化权限控制。账户锁定机制在5次失败登录后触发15分钟锁定。

2.3 训练引擎设计

训练引擎采用策略模式（Strategy Pattern）和模板方法模式（Template Method Pattern）设计。BaseTrainer抽象基类定义了训练流程模板（load_data→build_model→train_epoch×N→evaluate→save_model）。5个具体训练器分别为：

（1）sklearn_trainer：实现6个分类器（随机森林/SVM/KNN/逻辑回归/朴素贝叶斯/决策树）、6个回归器（线性回归/岭回归/Lasso/ElasticNet/SVR/KNN回归）和4个聚类器（KMeans/DBSCAN/层次聚类/GMM），支持warm_start增量训练和交叉验证。

（2）pytorch_trainer：MLP多层感知机，GPU加速（CUDA 12.4 + CuDNN），CosineAnnealingLR学习率调度，早停机制（patience=20 epoch），内置验证集划分和模型检查点保存。

（3）keras_trainer：TensorFlow/Keras Sequential API + BatchNorm，模型以.keras格式保存。

（4）transformers_nlp_trainer：HuggingFace Transformers库，支持BERT和DistilBERT中文文本分类微调，jieba分词 + 中文字符unigram混合分词。

（5）通用训练器：适用于ONNX及其他框架模型。

TrainingExecutor作为线程安全单例（双重检查锁定保证线程安全），基于ThreadPoolExecutor管理最多2个并发任务，通过threading.Event实现暂停/恢复/取消控制，配合EventBus和SSE流实现500ms级别的实时状态推送。

2.4 数据模型设计

系统核心数据模型包括四张主表。User表存储用户认证信息和角色权限（pbkdf2:sha256密码哈希）。Dataset表记录数据集文件路径、格式、类别、状态（uploading→processing→ready/error）和自动解析的统计摘要（summary_json），支持自引用测试集关联（source_dataset_id）。ModelRecord表管理模型元数据、5种框架类型、8种模型类型、多维度性能指标（accuracy/precision/recall/f1/R2/MSE/loss）和状态生命周期（draft→trained/deployed/archived），支持版本递增和独立测试集评估。TrainingJob表追踪训练任务完整执行过程，包括进度百分比、当前轮次/总轮次、资源需求（GPU/CPU/内存）、训练日志和每轮指标历史（metrics_history_json）。

2.5 前端技术选型

前端基于Bootstrap 5.3构建响应式布局，引入Vue 3 + Element Plus增强交互体验（el-backtop回到顶部、el-message消息提示、el-empty空状态、el-skeleton骨架屏、el-progress进度条、el-tag标签等组件），Chart.js实现训练过程中Accuracy/Loss曲线实时绘制。CSS采用GitHub风格暗色导航栏设计，支持亮/暗主题切换。Frontend v4.0新增Canvas粒子网络背景、滚动揭示动画、3D卡片倾斜效果和涟漪按钮点击反馈，beautify.css/js实现整体视觉美化。
"""

CH3_TITLE = '三、系统实施'
CH3_BODY = """
3.1 项目搭建与配置

项目基于Flask 3.1.0框架搭建，采用App Factory模式（app/__init__.py中的create_app函数）实现开发/测试/生产多环境配置切换，通过config.py管理不同环境的差异参数。SQLAlchemy 2.0 ORM统一数据访问层，配合Flask-Migrate（Alembic）管理数据库迁移脚本。Flask-Assets实现CSS/JS静态资源压缩打包，Flask-CORS 5.0.1配置跨域访问，flasgger 0.9.7+自动生成Swagger 3.0.3在线API文档。项目启动入口run.py仅需一条命令运行，访问http://127.0.0.1:5000进入系统。

3.2 用户认证实现

认证系统实现了三层防护机制。Web端使用Flask-Login管理Session，账户锁定在5次登录失败后触发15分钟锁定。JWT认证通过PyJWT签发，refresh_token配合token_version实现密码变更后令牌全局失效。API Key采用ak_前缀 + 48位十六进制随机串，SHA256哈希后存储。API兼容中间件自动将/api/*重写为/api/v1/*并添加废弃标头，确保渐进式升级。

3.3 数据集管理实现

数据集服务支持自动解析CSV/Excel/JSON/Parquet/TXT格式文件，检测缺失值、推断目标列、生成统计预览（row_count、column_count、缺失率、数据类型分布等）。数据集状态经历uploading→processing→ready/error完整生命周期。实现了15个公开数据源的自动导入（OpenML、UCI、Kaggle等），累计收集55个真实数据集，涵盖分类、回归、聚类、NLP、视觉、时序等任务类型。auto-config API会检查dataset.category并与算法匹配性发出警告。

3.4 训练执行引擎实现

训练引擎是系统的核心组件。TrainingExecutor作为线程安全单例，通过双重检查锁定（double-checked locking）保证线程安全，基于ThreadPoolExecutor管理并发任务。每个训练器在独立线程中运行，通过threading.Event机制实现暂停（pause_event.clear()）、恢复（pause_event.set()）和取消（cancel_event.set()）。

训练过程中，每个epoch结束时触发TrainingCallback回调函数，将当前进度百分比、损失值（loss）和评估指标写入TrainingJob的metrics_history_json字段，同时通过EventBus触发SSE事件。前端通过EventSource接口建立长连接，以500ms为间隔接收progress/complete/error/fatal事件，实现训练全过程的可视化监控。

系统目前支持5种训练框架。scikit-learn实现了6个分类器（随机森林、SVM、KNN、逻辑回归、朴素贝叶斯、决策树）、6个回归器（线性回归、岭回归、Lasso、ElasticNet、SVR、KNN回归）和4个聚类器（KMeans、DBSCAN、层次聚类、GMM），其中KMeans兼容sklearn 1.3+通过algorithm参数三重防护（lloyd/elvan/kmeans自动路由）解决冲突。PyTorch训练器实现MLP多层感知机，支持CUDA 12.4 GPU加速和CosineAnnealingLR学习率调度，内置早停机制和自动混合精度训练。TensorFlow/Keras训练器使用Sequential API + BatchNorm构建模型。Transformers NLP训练器支持BERT-base-chinese和DistilBERT中文文本分类微调，使用combined_tokenize()混合分词（jieba词分割 + 中文字符unigram）。所有训练器的random_state默认None实现真随机，传入int种子可复现。

3.5 超参数调优实现

超参数调优模块（hyperparameter_tuning.py，2124行）实现了三种搜索策略。GridSearchCV对参数组合穷举搜索并执行StratifiedKFold交叉验证，适用于小参数空间。RandomSearchCV对参数空间随机采样n_iter次，适用于高维参数空间。AutoML自动遍历所有适用算法，对每个算法执行快速RandomSearch，生成算法排名榜单。

聚类任务不支持交叉验证，系统通过_manual_clustering_search()方法使用ParameterGrid遍历 + 全量fit + 子采样silhouette_score评估。KMeans的algorithm参数冲突通过fix_kmeans_algorithm()三重防护解决。所有调优方法默认random_state=None实现真随机，传入int种子可复现结果。TuningProgressTracker单例以500ms间隔通过SSE推送实时进度。

3.6 模型质量治理

系统建立了完善的模型质量治理体系。quality_watchdog.py实现了once（一次性扫描）/daemon（持续守护）/quick（快速检查）三种模式，定期检测模型质量指标，发现过拟合或常数预测器后自动清理并通过SMTP/Webhook发送告警。diagnose_model_quality.py从7个维度诊断模型健康度——训练精度、验证精度、泛化差距、过拟合指数、常数预测检测、特征重要性分析、残差分析，按Tier 1（健康）/Tier 2（警告）/Tier 3（危险）分级管理。clean_overfit_models.py实现了批量安全删除和去重功能（--dedup参数）。

目前系统已训练138个模型（全部为已训练状态，0个草稿模型），其中sklearn模型133个、PyTorch模型5个，涵盖分类、回归、聚类和NLP任务。通过quality_watchdog清理了80个过拟合模型（clean_overfit_models.py执行）。质量分级结果：A级（健康）62个、B级（警告）46个、C级（危险）9个，过拟合和常数预测器已全部清理。系统通过375个pytest自动化测试用例验证，覆盖率达73%。

3.7 项目实施成果

项目累计产出Python应用代码60+文件，HTML模板36个，测试文件17个（375+测试用例），运维脚本48个。实现了60+ JSON API端点（7个Tag分组），Swagger 3.0.3在线文档自动生成。系统通过了完整的pytest测试套件，覆盖认证、数据集、模型、训练等核心功能模块，质量看门狗持续运行守护模型健康。
"""

CH4_TITLE = '四、总结'
CH4_BODY = """
4.1 项目成果

通过本次专业实训，成功完成了一个功能完整的AI模型训练管理平台"myfirstaiproject"的开发和部署。系统实现了从数据集管理、模型训练、超参数调优到模型推理的完整MLOps工作流，支持scikit-learn、PyTorch、TensorFlow/Keras、Transformers四种主流机器学习框架，累计训练138个真实模型，管理55个真实数据集，通过375个自动化测试用例验证，整体代码覆盖率达73%。

4.2 关键技术难点

（1）多框架训练器统一抽象：不同训练框架的训练流程、数据格式和模型接口差异巨大，通过BaseTrainer抽象基类和策略模式实现了统一的训练接口。模板方法模式定义了标准训练流程（load_data→build_model→train→evaluate→save），各训练器只需实现各自框架特定的逻辑。

（2）实时训练监控：采用SSE（Server-Sent Events）替代WebSocket，利用HTTP长连接和事件流机制实现500ms级别的实时进度更新。EventBus解耦训练引擎和SSE推送，支持多订阅者模式。

（3）超参数调优与聚类冲突：聚类算法无监督特性与交叉验证不兼容，通过_manual_clustering_search()方法实现自适应的调优策略。KMeans的algorithm参数在不同scikit-learn版本中的兼容性问题通过fix_kmeans_algorithm()三重防护机制解决。

（4）Windows GBK环境兼容：在Windows环境下处理中文文本时，通过显式指定UTF-8编码、避免emoji字符和统一字符串处理方式确保系统稳定性。TF-IDF特征保护机制确保StandardScaler不应用于tfidf_*列。

4.3 改进方向

未来可以从以下方向继续优化：引入Redis缓存加速频繁查询和大规模模型加载；支持ONNX Runtime实现跨平台模型部署和优化；增加分布式训练支持以应对更大规模的数据集和模型；完善CI/CD流水线和GitOps自动化部署流程；扩展MCP协议集成实现LLM驱动的智能模型推荐。
"""

HARVEST_TITLE = '实训收获与体会'
HARVEST_BODY = """
通过本次为期三周的专业实训，我在Web全栈开发、机器学习工程化和AI系统架构设计方面获得了丰富的实践经验，受益良多。

一、技术能力提升

在前端开发方面，我深入掌握了Jinja2模板引擎与Vue 3、Element Plus、Bootstrap 5.3的混合使用技巧，学会了如何构建响应式、现代化的数据管理界面。Chart.js的集成让我能够实时可视化训练过程中的Accuracy和Loss曲线变化，直观地理解模型收敛过程。Canvas粒子网络和3D卡片倾斜等前端v4.0特效的实现，提升了对Web前端美学的理解。

在后端开发方面，我系统学习了Flask 3.1的应用工厂模式、Blueprint模块化路由组织和SQLAlchemy 2.0 ORM的最佳实践（db.session.get()替代旧版query API）。通过实现三位一体认证架构（Session + JWT + API Key），我加深了对Web安全机制和token管理的理解。Service层与Route层的职责分离让我体会到了分层架构在大规模项目中的维护优势。

在机器学习工程化方面，我掌握了从数据集加载、预处理（SMOTE过采样/欠采样）、训练到评估的完整流程。通过实现BaseTrainer抽象基类和5个具体训练器，我深入理解了策略模式和模板方法模式在实际项目中的应用。GPU加速的PyTorch训练器实现让我体会到了深度学习框架的底层工作机制和CUDA/CuDNN加速原理。

二、工程实践能力

项目开发过程中，我面对了多个真实工程挑战。解决KMeans algorithm参数兼容性问题时，我学会了在多版本库环境下编写防御性代码（兼容sklearn 1.3+/1.6+）。实现SSE实时推送时，我深入理解了HTTP长连接和事件流机制，以及EventBus发布订阅模式在解耦系统组件中的作用。处理Windows GBK中文编码问题时，我积累了跨平台兼容性调试经验。TF-IDF特征保护的设计让我认识到数据预处理流水线中操作顺序的重要性。

项目累计完成375个自动化测试用例，代码覆盖率达到73%，这种测试先行的开发方式保证了系统在持续迭代中的稳定性。质量治理体系（quality_watchdog + diagnose_model_quality + clean_overfit_models）的实现让我认识到，一个成熟的AI平台不仅需要关注模型训练，还需要建立完善的模型质量监控和自动清理机制。

三、团队协作与项目管理

实训过程中，我遵循了规范的Git工作流程，每次修改代码后自动提交并附上描述性信息。通过CLAUDE.md文件维护项目约定文档，记录架构决策、代码规范和重要约定。编写ARCHITECTURE.md和PRD.md文档的过程，让我学会了如何从架构层面思考和设计系统，以及如何通过文档传递设计意图。

四、总结

本次实训不仅提升了我的技术能力，更重要的是培养了我从系统架构角度思考问题的能力。从数据集管理到大模型训练，从单一算法到多框架集成，从手工调参到自动搜索，我深刻体会到了MLOps工程化的核心价值——让AI开发更高效、更规范、更可控。这段经历将为我未来在人工智能领域的深入发展打下坚实基础。
"""


# ──────────────────────────────────────────────
#  Helper functions
# ──────────────────────────────────────────────

def _make_run(text: str, font_name: str = '宋体', font_size: Pt = Pt(12),
              bold: bool = False, color: RGBColor | None = None) -> dict:
    """Build a run descriptor dict for convenience."""
    return {
        'text': text,
        'font_name': font_name,
        'font_size': font_size,
        'bold': bold,
        'color': color,
    }


def _set_run_font(run, font_name: str, font_size: Pt, bold: bool,
                  color: RGBColor | None = None) -> None:
    """Apply font properties to a run."""
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_paragraph_to_cell(cell, text: str, font_name: str = '宋体',
                          font_size: Pt = Pt(12), bold: bool = False,
                          alignment: int = WD_ALIGN_PARAGRAPH.LEFT,
                          first_line_indent: Cm | None = None,
                          line_spacing: float = 1.5,
                          space_before: Pt = Pt(0), space_after: Pt = Pt(2)):
    """Append a formatted paragraph to a table cell."""
    p = cell.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    run = p.add_run(text)
    _set_run_font(run, font_name, font_size, bold)
    return p


def _fill_cell_with_content(cell, title: str,
                            chapters: list[tuple[str, str]]) -> None:
    """Clear a table cell and fill it with structured chapter content.

    Args:
        cell: The table cell to fill.
        title: Row-level title (e.g. "实训内容").
        chapters: List of (chapter_title, chapter_body) pairs.
    """
    # Clear existing paragraphs
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)

    # Row title
    add_paragraph_to_cell(cell, title, font_name='宋体', font_size=Pt(14),
                          bold=True, space_before=Pt(4))

    for ch_title, ch_body in chapters:
        # Chapter spacing
        add_paragraph_to_cell(cell, '', font_size=Pt(6), space_after=Pt(0))
        # Chapter title
        add_paragraph_to_cell(cell, ch_title, font_name='黑体', font_size=Pt(16),
                              bold=True, space_before=Pt(6), space_after=Pt(4))

        sections = [s.strip() for s in ch_body.strip().split('\n\n') if s.strip()]
        for section in sections:
            for line in section.split('\n'):
                line = line.strip()
                if not line:
                    continue
                is_subheading = (
                    line and len(line) > 2
                    and line[0].isdigit()
                    and any(c in line[:5] for c in ('.', '、', '）'))
                )
                if is_subheading:
                    add_paragraph_to_cell(cell, line, font_name='黑体',
                                          font_size=Pt(12), bold=True,
                                          space_after=Pt(1))
                else:
                    add_paragraph_to_cell(cell, line, font_name='宋体',
                                          font_size=Pt(12),
                                          first_line_indent=Cm(0.74),
                                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def _fill_cell_with_harvest(cell, title: str, body: str) -> None:
    """Fill the harvest & reflection cell."""
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)

    add_paragraph_to_cell(cell, title, font_name='宋体', font_size=Pt(14),
                          bold=True, space_before=Pt(4))

    sections = [s.strip() for s in body.strip().split('\n\n') if s.strip()]
    for section in sections:
        lines = section.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            is_subheading = (
                line and len(line) > 2
                and line[0].isdigit()
                and '、' in line[:3]
            )
            if is_subheading:
                add_paragraph_to_cell(cell, line, font_name='黑体',
                                      font_size=Pt(12), bold=True)
            else:
                add_paragraph_to_cell(cell, line, font_name='宋体',
                                      font_size=Pt(12),
                                      first_line_indent=Cm(0.74),
                                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def _center_cell_vertically(cell) -> None:
    """Set vertical alignment of a table cell to center."""
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tc_pr.append(v_align)


# ──────────────────────────────────────────────
#  Main generation
# ──────────────────────────────────────────────

def generate_report(student: StudentInfo, cfg: FileConfig) -> Path:
    """Generate the full training report docx.

    Args:
        student: Student personal information.
        cfg: File configuration.

    Returns:
        Path to the generated docx file.

    Raises:
        FileNotFoundError: If a template file does not exist.
    """
    # Validate template files exist
    for name, fpath in [('Report template', cfg.template_report),
                        ('Assessment form', cfg.template_form)]:
        if not fpath.is_file():
            raise FileNotFoundError(f'{name} not found: {fpath}')

    # 1. Open report template
    doc = Document(str(cfg.template_report))

    # 2. Fill cover info table (Table 0)
    info_table = doc.tables[0]
    info_data = [
        (1, '《专业实训(web软件开发方向)》'),
        (1, f'{student.year}年6月16日起'),
        (1, f'{student.year}年7月04日止'),
        (1, student.name),
        (1, student.student_id),
        (1, student.class_name),
        (1, student.teacher),
    ]
    for idx, (col_idx, text) in enumerate(info_data):
        cell = info_table.rows[idx].cells[col_idx]
        # Clear default paragraph
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        _set_run_font(run, '宋体', Pt(14), bold=(idx <= 2))
        _center_cell_vertically(cell)

    # 3. Fill content body (Table 1)
    content_table = doc.tables[1]

    # Row 0: 实训内容 (chapters 1-4)
    chapters = [
        (CH1_TITLE, CH1_BODY),
        (CH2_TITLE, CH2_BODY),
        (CH3_TITLE, CH3_BODY),
        (CH4_TITLE, CH4_BODY),
    ]
    _fill_cell_with_content(content_table.rows[0].cells[0],
                            '实训内容', chapters)

    # Row 1: 实训收获与体会
    _fill_cell_with_harvest(content_table.rows[1].cells[0],
                            '实训收获与体会', HARVEST_BODY)

    # Row 2: 实训成绩 (keep template content as-is)

    # 4. Append assessment form via XML cloning
    doc.add_page_break()

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('《专业综合实训II(校企合作)》考核表')
    _set_run_font(run, '宋体', Pt(16), bold=True)
    title_p.paragraph_format.space_after = Pt(6)

    # Info line
    info_p = doc.add_paragraph()
    info_run = info_p.add_run(f'姓名：{student.name}\t\t学号：{student.student_id}')
    _set_run_font(info_run, '宋体', Pt(12), bold=False)
    info_p.paragraph_format.space_after = Pt(6)

    # Clone table XML from assessment form
    form_doc = Document(str(cfg.template_form))
    for src_table in form_doc.tables:
        tbl_clone = copy.deepcopy(src_table._tbl)

        # Set table width to 100%
        ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        tbl_pr = tbl_clone.find(f'{{{ns_w}}}tblPr')
        if tbl_pr is None:
            tbl_pr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl_clone.insert(0, tbl_pr)
        existing = tbl_pr.find(f'{{{ns_w}}}tblW')
        if existing is not None:
            tbl_pr.remove(existing)
        tbl_w = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
        tbl_pr.insert(0, tbl_w)

        doc.element.body.append(tbl_clone)
        doc.add_paragraph()

    # Note
    note_p = doc.add_paragraph()
    note_run = note_p.add_run('说明：本考核表一式两份，一份附在课程报告后面，一份由学院存档。')
    _set_run_font(note_run, '宋体', Pt(10), bold=False, color=RGBColor(0x66, 0x66, 0x66))
    note_p.paragraph_format.space_before = Pt(6)

    # 5. Save
    doc.save(str(cfg.output_path))
    return cfg.output_path


def main() -> None:
    """Entry point: generate and verify the report."""
    cfg = FileConfig()
    student = StudentInfo()

    print('=' * 60)
    print('  Generating training report...')
    print('=' * 60)

    try:
        out = generate_report(student, cfg)
        print(f'  Done: {out}')
        print(f'  Size: {out.stat().st_size:,} bytes')
    except FileNotFoundError as e:
        print(f'  ERROR: {e}')
        raise
    except Exception as e:
        print(f'  ERROR: {e}')
        raise

    print('=' * 60)


if __name__ == '__main__':
    main()
