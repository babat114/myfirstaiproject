#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Verify the generated report"""

import sys
import os

from docx import Document

path = r'C:\Users\86180\Desktop\校内专业实训(web软件开发方向)报告 - 李想.docx'
doc = Document(path)

print('=== Report Verification ===')
print()

# 1. Cover info
t0 = doc.tables[0]
print('[1] Cover Info Table:')
labels = ['课程名称', '实习时间起', '实习时间止', '学生姓名', '学号', '专业班级', '指导教师']
values = []
for ri in range(7):
    v = t0.rows[ri].cells[1].text.strip()
    values.append(v)
    l = labels[ri] if ri < len(labels) else f'Row{ri}'
    print(f'    {l}: {v}')

assert '李想' in values[3], f'Name not found: {values[3]}'
assert '2023081213' in values[4], f'ID not found: {values[4]}'
assert '计科232' in values[5], f'Class not found: {values[5]}'
assert '赵迪' in values[6], f'Teacher not found: {values[6]}'
print('  >> Cover info: ALL CORRECT')
print()

# 2. Content table
t1 = doc.tables[1]
print('[2] Content Table:')
row0_text = t1.rows[0].cells[0].text
row1_text = t1.rows[1].cells[0].text

chapters = ['需求分析', '系统设计', '系统实施', '总结']
for ch in chapters:
    found = ch in row0_text
    print(f'    Chapter {ch}: {"PASS" if found else "FAIL"}')
    assert found, f'Missing chapter: {ch}'

assert '收获' in row1_text, 'Missing 收获体会'
print('    Harvest section: PASS')
print(f'    Content chars: {len(row0_text)} (req >= 4000)')
assert len(row0_text) >= 4000, f'Content too short: {len(row0_text)}'
print('  >> Content: ALL PASS')
print()

# 3. Assessment table
t2 = doc.tables[2]
print('[3] Assessment Table:')
assert len(t2.rows) == 13, f'Expected 13 rows, got {len(t2.rows)}'
assert len(t2.columns) == 6, f'Expected 6 cols, got {len(t2.columns)}'
print(f'    Rows: {len(t2.rows)}, Cols: {len(t2.columns)}')

# Verify all rows have correct number of logical cells
for ri in range(13):
    nc = len(t2.rows[ri].cells)
    print(f'    Row {ri}: {nc} cells')
    assert nc == 6, f'Row {ri} has {nc} cells instead of 6'

print('  >> Assessment table: ALL PASS')
print()

# 4. Total stats
total = sum(len(p.text) for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            total += len(c.text)
print(f'[4] Stats:')
print(f'    Total characters: {total}')
print(f'    File size: {os.path.getsize(path):,} bytes')
print(f'    sections: {len(doc.sections)}')
print(f'    paragraphs: {len(doc.paragraphs)}')
print(f'    tables: {len(doc.tables)}')
print()

print('=' * 50)
print('ALL CHECKS PASSED - Report is valid!')
print('=' * 50)
